from pathlib import Path
from docx import Document


def read_docx(file_path: Path) -> str:
    # open the word document
    doc = Document(file_path)

    parts: list[str] = []

    # extract text from paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    # extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells: list[str] = []

            for cell in row.cells:
                value = cell.text.strip()
                if value:
                    cells.append(value)

            if cells:
                parts.append(" | ".join(cells))

    # join all extracted content into a single transcript string
    return "\n".join(parts)