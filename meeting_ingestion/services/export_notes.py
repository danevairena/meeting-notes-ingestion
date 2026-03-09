import re
from pathlib import Path

from docx import Document

from ..supabase_client import get_supabase_client


# sanitize project folder names so they are safe for filesystem usage
def _safe_folder_name(name: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*]', "", name or "")
    cleaned = cleaned.strip().rstrip(".")
    return cleaned or "unknown-project"


# load all meetings that already have notes
def _get_meetings_with_notes() -> list[dict]:
    supabase = get_supabase_client()

    result = (
        supabase.table("meetings")
        .select(
            """
            id,
            title,
            meeting_date,
            source_file,
            projects(name),
            notes(summary,action_items,key_takeaways,topics,next_steps)
            """
        )
        .execute()
    )

    rows = result.data or []
    meetings_with_notes = []

    # normalize supabase nested response and attach project name
    for row in rows:
        notes = row.get("notes")
        if not notes:
            continue

        if isinstance(notes, list):
            if not notes:
                continue
            notes = notes[0]

        project = row.get("projects")

        if isinstance(project, list):
            project = project[0] if project else None

        project_name = "unknown-project"
        if isinstance(project, dict):
            project_name = project.get("name") or project_name

        row["notes"] = notes
        row["project_name"] = _safe_folder_name(project_name)

        meetings_with_notes.append(row)

    return meetings_with_notes


# export one meeting notes record to a docx file
def _export_single_meeting(meeting: dict, output_dir: Path) -> Path:
    notes = meeting["notes"]
    title = meeting.get("title") or "Untitled meeting"
    meeting_date = meeting.get("meeting_date") or "unknown-date"
    project_name = meeting.get("project_name") or "unknown-project"

    # create project directory inside export folder
    project_dir = output_dir / project_name
    project_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Meeting ID: {meeting['id']}")
    doc.add_paragraph(f"Meeting Date: {meeting_date}")
    doc.add_paragraph(f"Project: {project_name}")

    # summary section
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(notes.get("summary") or "")

    # key takeaways section
    doc.add_heading("Key Takeaways", level=2)
    for item in notes.get("key_takeaways") or []:
        doc.add_paragraph(str(item), style="List Bullet")

    # topics section
    doc.add_heading("Topics", level=2)
    for item in notes.get("topics") or []:
        doc.add_paragraph(str(item), style="List Bullet")

    # action items section
    doc.add_heading("Action Items", level=2)

    action_items = notes.get("action_items") or []

    if action_items:
        for item in action_items:
            text = item.get("text") or ""
            if not text:
                continue

            owner = item.get("owner")
            due_date = item.get("due_date")

            # build optional metadata for owner and due date
            details: list[str] = []
            if owner:
                details.append(f"owner: {owner}")
            if due_date:
                details.append(f"due: {due_date}")

            line = text
            if details:
                line = f"{text} — {', '.join(details)}"

            doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_paragraph("No action items detected.")

    # next steps section
    doc.add_heading("Next Steps", level=2)

    next_steps = notes.get("next_steps") or []

    if next_steps:
        for item in next_steps:
            text = item.get("text") or ""
            if not text:
                continue

            owner = item.get("owner")

            line = text
            if owner:
                line = f"{text} — owner: {owner}"

            doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_paragraph("No next steps detected.")

    # reuse original transcript filename for export
    original_file_name = Path(
        meeting.get("source_file") or f"{meeting['id']}.docx"
    ).name

    file_path = project_dir / original_file_name

    # skip export if file already exists
    if file_path.exists():
        return file_path

    doc.save(file_path)

    return file_path


# export all meetings with notes to docx files
def export_all_notes(output_dir: str = "exports") -> list[Path]:
    export_path = Path(output_dir)
    export_path.mkdir(parents=True, exist_ok=True)

    meetings = _get_meetings_with_notes()

    exported_files: list[Path] = []

    # export each meeting notes document
    for meeting in meetings:
        exported_file = _export_single_meeting(meeting, export_path)
        exported_files.append(exported_file)

    return exported_files